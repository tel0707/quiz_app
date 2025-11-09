# quiz/views.py
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.models import User
from django.db.models import Prefetch, Case, When
from .forms import QuizTypeForm, QuestionForm, AnswerFormSet, AnswerUpdateFormSet, UploadWordForm
from django.views.generic import CreateView, UpdateView, DeleteView, ListView
from django.urls import reverse_lazy
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from django.contrib.messages.views import SuccessMessageMixin
from docx import Document
import re
import random, json
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.core.paginator import Paginator
from django.http import JsonResponse
from django.db import transaction
from django.utils import timezone
from .models import QuizType, Question, Answer, GenerateQuiz, AnswerUsers, GenerateQuizQuestion



# --- Ro‘yxat (hamma ko‘ra oladi) ---
class QuizTypeListView(ListView):
    model = QuizType
    template_name = 'quiz/quiztype_list.html'
    context_object_name = 'quiztypes'
    paginate_by = 10

    def get_queryset(self):
        return QuizType.objects.filter(is_active=True).order_by('-id')


# --- Superuser tekshiruvi ---
class SuperuserRequiredMixin(LoginRequiredMixin, UserPassesTestMixin):
    def test_func(self):
        return self.request.user.is_superuser

    def handle_no_permission(self):
        messages.error(self.request, "Bu amalni faqat administrator bajarishi mumkin.")
        return super().handle_no_permission()


# --- Yaratish (faqat superuser) ---
class QuizTypeCreateView(SuperuserRequiredMixin, SuccessMessageMixin, CreateView):
    model = QuizType
    form_class = QuizTypeForm
    template_name = 'quiz/quiztype_form.html'
    success_url = reverse_lazy('quiztype_list')
    success_message = "Quiz turi muvaffaqiyatli qo‘shildi!"


# --- Tahrirlash (faqat superuser) ---
class QuizTypeUpdateView(SuperuserRequiredMixin, SuccessMessageMixin, UpdateView):
    model = QuizType
    form_class = QuizTypeForm
    template_name = 'quiz/quiztype_form.html'
    success_url = reverse_lazy('quiztype_list')
    success_message = "Quiz turi muvaffaqiyatli yangilandi!"



    def form_valid(self, form):
        context = self.get_context_data()
        formset = context['formset']

        if not formset.is_valid():
            return self.form_invalid(form)

        # Yangi va o'chirilmagan to'g'ri javoblarni hisoblaymiz
        correct_answers = [
            f for f in formset.forms
            if f.cleaned_data and not f.cleaned_data.get('DELETE', False) and f.cleaned_data.get('is_correct')
        ]

        if len(correct_answers) == 0:
            messages.error(self.request, "Kamida bitta to'g'ri javob bo'lishi kerak!")
            return self.form_invalid(form)

        with transaction.atomic():
            self.object = form.save()
            formset.instance = self.object
            formset.save()

            quiz_type_id = form.cleaned_data['quiz_type'].id
            self.request.session['last_quiztype_id'] = quiz_type_id

        return super().form_valid(form)


# --- O‘chirish (faqat superuser) ---
class QuizTypeDeleteView(SuperuserRequiredMixin, SuccessMessageMixin, DeleteView):
    model = QuizType
    template_name = 'quiz/quiztype_confirm_delete.html'
    success_url = reverse_lazy('quiztype_list')
    success_message = "Quiz turi muvaffaqiyatli o‘chirildi!"

    def delete(self, request, *args, **kwargs):
        messages.success(self.request, self.success_message)
        return super().delete(request, *args, **kwargs)


# Superuser tekshiruvi
class SuperuserRequiredMixin(LoginRequiredMixin, UserPassesTestMixin):
    def test_func(self):
        return self.request.user.is_superuser

    def handle_no_permission(self):
        messages.error(self.request, "Bu amalni faqat administrator bajarishi mumkin.")
        return redirect('question_list')


# --- Savollar ro'yxati ---
class QuestionListView(ListView):
    model = Question
    template_name = 'quiz/question_list.html'
    context_object_name = 'questions'
    paginate_by = 10

    def get_queryset(self):
        return Question.objects.select_related('quiz_type').prefetch_related('answers').filter(is_active=True).order_by('-id')


# --- Savol qo'shish ---
class QuestionCreateView(SuperuserRequiredMixin, SuccessMessageMixin, CreateView):
    model = Question
    form_class = QuestionForm
    template_name = 'quiz/question_form.html'
    success_url = reverse_lazy('question_list')
    success_message = "Savol va javoblar muvaffaqiyatli qo'shildi!"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        if self.request.POST:
            context['formset'] = AnswerFormSet(self.request.POST)
        else:
            context['formset'] = AnswerFormSet()
        return context

    def form_valid(self, form):
        context = self.get_context_data()
        formset = context['formset']

        if formset.is_valid():
            correct_answers = [
                f for f in formset.forms
                if f.cleaned_data and not f.cleaned_data.get('DELETE', False) and f.cleaned_data.get('is_correct')
            ]

            if len(correct_answers) == 0:
                messages.error(self.request, "Kamida bitta to'g'ri javob bo'lishi kerak!")
                return self.form_invalid(form)
            self.object = form.save()
            formset.instance = self.object
            formset.save()
            # Oxirgi tanlangan quiz_type ni saqlash
            quiz_type_id = form.cleaned_data['quiz_type'].id
            self.request.session['last_quiztype_id'] = quiz_type_id
            return super().form_valid(form)
        else:
            return self.form_invalid(form)


# --- Savol tahrirlash ---
class QuestionUpdateView(SuperuserRequiredMixin, SuccessMessageMixin, UpdateView):
    model = Question
    form_class = QuestionForm
    template_name = 'quiz/question_form.html'
    success_url = reverse_lazy('question_list')
    success_message = "Savol muvaffaqiyatli yangilandi!"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        if self.request.POST:
            context['formset'] = AnswerUpdateFormSet(self.request.POST, instance=self.object)
        else:
            context['formset'] = AnswerUpdateFormSet(instance=self.object)
        return context

    def form_valid(self, form):
        context = self.get_context_data()
        formset = context['formset']
        if formset.is_valid():
            self.object = form.save()
            formset.instance = self.object
            formset.save()
            # Yangi tanlangan quiz_type ni saqlash
            quiz_type_id = form.cleaned_data['quiz_type'].id
            self.request.session['last_quiztype_id'] = quiz_type_id
            return super().form_valid(form)
        else:
            return self.form_invalid(form)


# --- Savol o'chirish ---
class QuestionDeleteView(SuperuserRequiredMixin, SuccessMessageMixin, DeleteView):
    model = Question
    template_name = 'quiz/question_confirm_delete.html'
    success_url = reverse_lazy('question_list')
    success_message = "Savol muvaffaqiyatli o‘chirildi!"

    def form_valid(self, form):
        # SuccessMessageMixin avtomatik ishlaydi, lekin agar qo'shimcha logika kerak bo'lsa:
        messages.success(self.request, self.success_message)
        return super().form_valid(form)


def upload_quiz_from_word(request):
    """
    Word fayldan savollarni import qiluvchi view.
    - Hujjat ichidagi PARAGRAFLAR va TABLE CELL-larni ketma-ket o'qiydi.
    - '1. ' kabi raqam+nuqta bilan boshlangan qatorlarni yangi savol deb oladi.
    - Raqamli qator va keyingi kelgan barcha qatorlar o'rtasidagi satrlar javob variantlari hisoblanadi.
    - Javob satri boshida '*' bo'lsa -> is_correct = True.
    """
    if request.method == "POST":
        form = UploadWordForm(request.POST, request.FILES)
        if form.is_valid():
            quiz_type = form.cleaned_data['quiz_type']
            file = request.FILES['file']

            try:
                doc = Document(file)
            except Exception as e:
                messages.error(request, f"❌ Word faylni o‘qib bo‘lmadi: {e}")
                return redirect("upload_quiz_from_word")

            # --- 1) Hujjatdan barcha matn qatorlarini yig'ish ---
            lines = []

            # Paragraflar (agar oddiy matn ham bo'lsa)
            for p in doc.paragraphs:
                t = p.text.strip()
                if t:
                    # no-break space to'g'irlash
                    lines.append(t.replace("\xa0", " "))

            # Jadval ichidagi har bir cellni ham ketma-ket qo'shamiz
            for table in doc.tables:
                for row in table.rows:
                    # Agar jadvalda bir nechta cell bo'lsa, ularni birlashtiramiz (odatiy holat: 1 cell/row)
                    # Ammo bu yerda har bir cell alohida satr bo'lib keladi: shuning uchun hamma celllarni alohida qo'shamiz
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            lines.append(cell_text.replace("\xa0", " "))

            if not lines:
                messages.error(request, "❌ Fayldan hech qanday matn o‘qilmadi.")
                return redirect("upload_quiz_from_word")

            # --- 2) Lines ni tahlil qilib savol-javob bloklariga ajratish ---
            question_blocks = []  # har element: (q_line, [answer_lines...])
            current_q = None
            current_answers = []

            # regex: boshida raqam + nuqta (masalan "1. " yoki "12. ")
            q_pattern = re.compile(r'^\s*(\d+)\.\s*(.*)')

            for line in lines:
                # Tozalash
                line = line.strip()
                if not line:
                    continue

                m = q_pattern.match(line)
                if m:
                    # yangi savol topildi
                    # avvalgi savolni saqlaymiz (agar mavjud)
                    if current_q is not None:
                        question_blocks.append((current_q, current_answers))
                    # boshlagich: raqam va qolgan matn (savol matni)
                    num = m.group(1)
                    rest = m.group(2).strip()
                    # Agar savol matn bo'sh bo'lsa, butun line ni savol sifatida olamiz
                    if not rest:
                        rest = line
                    current_q = f"{num}. {rest}"
                    current_answers = []
                else:
                    # Bu qatorda javob variantlari yoki savolga tegishli davomiy qatorlar bor
                    # Agar qatorda bir nechta javoblar bo'lib ketgan bo'lsa (kam ehtimol), biz har bir qatorni alohida javob sifatida qabul qilamiz.
                    current_answers.append(line)

            # Oxirgi savolni ham qo'shamiz
            if current_q is not None:
                question_blocks.append((current_q, current_answers))

            if not question_blocks:
                messages.error(request, "❌ Faylda savollar topilmadi. Iltimos formatni tekshiring (1. Savol ...).")
                return redirect("upload_quiz_from_word")

            # --- 3) Har bir blokni DB ga yozish ---
            created_q = 0
            problems = []
            for q_text, answers in question_blocks:
                # Ba'zi bloklar savol bilan javobsiz kelishi mumkin -> e'tibor
                if not answers:
                    problems.append(f"'{q_text}' uchun javob topilmadi")
                    continue

                # Savolni saqlaymiz
                question = Question.objects.create(
                    quiz_type=quiz_type,
                    name=q_text,
                )

                # Har bir javob qatorini Answer qilib yozamiz.
                for ans_line in answers:
                    # to'g'ri javob belgisi: boshida '*' bo'lsa yoki ans_line ichida '*'
                    is_correct = False
                    text = ans_line
                    # ba'zan '*' belgisi qatorda boshlang‘ichda, ba'zan oxirida — oddiy holatda boshlanishi
                    if text.startswith("*"):
                        is_correct = True
                        text = text[1:].strip()
                    else:
                        # agar '*' boshqa joyda bo'lsa (kam uchraydi), olib tashlab flag qo'yamiz
                        if "*" in text:
                            is_correct = True
                            text = text.replace("*", "").strip()

                    # Oxirgi tozalash
                    text = text.strip()

                    # Agar javob bo'sh bo'lsa — o'tkazamiz
                    if not text:
                        continue

                    Answer.objects.create(
                        question=question,
                        name=text,
                        is_correct=is_correct
                    )

                created_q += 1

            # --- 4) Xabar berish ---
            msg = f"✅ {created_q} ta savol '{quiz_type.name}' turiga muvaffaqiyatli yuklandi."
            if problems:
                msg += " Ba'zi bloklarda muammo: " + "; ".join(problems[:5])
            messages.success(request, msg)
            return redirect("upload_quiz_from_word")

        else:
            messages.error(request, "❌ Forma to'ldirishda xatolik bor.")
            return redirect("upload_quiz_from_word")
    else:
        form = UploadWordForm()

    return render(request, "quiz/upload_quiz.html", {"form": form})




# 🔹 TESTNI BOSHLASH
# quiz/views.py da generate_quiz ni shu bilan almashtiring
def generate_quiz(request, pk):
    n = int(request.GET.get("count", 10))
    quiz_type = get_object_or_404(QuizType, pk=pk)

    # Faol savollar
    question_ids = list(
        Question.objects.filter(is_active=True, quiz_type_id=pk)
        .values_list('id', flat=True)
    )
    if not question_ids:
        messages.error(request, f"❌ '{quiz_type.name}' uchun faol savollar yo‘q.")
        return redirect('quiztype_list')

    n = min(n, len(question_ids))
    selected_q_ids = random.sample(question_ids, n)

    with transaction.atomic():
        quiz = GenerateQuiz.objects.create(
            user=request.user,
            quiz_type=quiz_type
        )

        # MUHIM: SESSIYAGA TO‘G‘RI MA'LUMOTLAR
        request.session['quiz_id'] = quiz.id
        request.session['selected_q_ids'] = selected_q_ids
        request.session['quiz_start_time'] = timezone.now().isoformat()  # TIMER UCHUN
        request.session.modified = True  # Django sessiyani saqlasin

        # Savollarni GenerateQuizQuestion ga yozamiz
        for qid in selected_q_ids:
            GenerateQuizQuestion.objects.create(
                quiz=quiz,
                question_id=qid
            )

    # TO‘G‘RI yo‘nalish: quiz_id bilan
    return redirect('quiz_page', quiz_id=quiz.id, page=1)

# 🔹 SAVOLLARNI KO‘RISH (har biri alohida sahifada)
def quiz_page(request, quiz_id, page):
    quiz = get_object_or_404(GenerateQuiz, id=quiz_id, user=request.user)
    selected_q_ids = request.session.get('selected_q_ids', [])

    # Filter active questions, preserving session order
    questions = Question.objects.filter(id__in=selected_q_ids, is_active=True).prefetch_related(
        Prefetch('answers', queryset=Answer.objects.filter(is_active=True))
    )

    if selected_q_ids:
        preserved_order = Case(*[When(id=pk, then=pos) for pos, pk in enumerate(selected_q_ids)])
        questions = questions.order_by(preserved_order)

    paginator = Paginator(questions, 1)
    page_obj = paginator.get_page(page)
    question = page_obj.object_list[0] if page_obj else None

    # JAVOB BERILGAN SAHIFA RAQAMLARINI hisoblaymiz
    answered_question_ids = AnswerUsers.objects.filter(
        user=request.user,
        generate_quiz=quiz
    ).values_list('question_id', flat=True)

    # Har bir savolning tartib raqamini (1,2,3...) topamiz
    answered_pages = []
    for idx, q in enumerate(questions, 1):
        if q.id in answered_question_ids:
            answered_pages.append(idx)

    context = {
        'quiz_type': quiz.quiz_type,
        'quiz': quiz,
        'question': question,
        'paginator': paginator,
        'page_obj': page_obj,
        'answered_questions': answered_pages,  # SAHIFA RAQAMLARI (1,2,3...)
        'total_minutes': len(selected_q_ids),
    }
    return render(request, 'quiz/quiz_page.html', context)

# 🔹 AJAX ORQALI JAVOBNI SAQLASH
# quiz/views.py
def save_answer(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            q_id = int(data.get('question_id'))
            a_id = int(data.get('answer_id'))
            quiz_id = request.session.get('quiz_id')

            if not quiz_id:
                return JsonResponse({"success": False, "error": "Sessiya yo'qolgan"}, status=400)

            quiz = get_object_or_404(GenerateQuiz, id=quiz_id, user=request.user)
            question = get_object_or_404(Question, id=q_id)
            answer = get_object_or_404(Answer, id=a_id)

            # Eski javobni o'chirish
            AnswerUsers.objects.filter(
                user=request.user,
                generate_quiz=quiz,
                question=question
            ).delete()

            # Yangi javob saqlash
            AnswerUsers.objects.create(
                user=request.user,
                generate_quiz=quiz,
                question=question,
                answer=answer
            )

            # Javob berilgan savollar ro'yxatini yangilash
            answered = AnswerUsers.objects.filter(
                user=request.user,
                generate_quiz=quiz
            ).values_list('question_id', flat=True)

            selected_q_ids = request.session.get('selected_q_ids', [])
            answered = set(
                AnswerUsers.objects.filter(user=request.user, generate_quiz=quiz).values_list('question_id', flat=True))
            answered_pages = [idx for idx, qid in enumerate(selected_q_ids, 1) if qid in answered]
            return JsonResponse({
                "success": True,
                "answered_questions": answered_pages
            })

        except Exception as e:
            return JsonResponse({"success": False, "error": str(e)}, status=500)

    return JsonResponse({"success": False}, status=400)


# 🔹 TESTNI TUGATISH
def finish_quiz(request):
    quiz_id = request.session.get('quiz_id')
    if not quiz_id:
        return redirect('quiztype_list')

    quiz = get_object_or_404(GenerateQuiz, id=quiz_id, user=request.user)
    selected_q_ids = request.session.get('selected_q_ids', [])
    total_questions = len(selected_q_ids)

    answers = AnswerUsers.objects.filter(user=request.user, generate_quiz=quiz)
    correct = sum(1 for ans in answers if ans.answer.is_correct)

    quiz.score = int((correct / total_questions) * 100) if total_questions > 0 else 0
    quiz.finished = timezone.now()
    quiz.save()

    # Sessiyani tozalaymiz
    for key in ['quiz_id', 'selected_q_ids', 'quiz_start_time']:
        request.session.pop(key, None)

    return render(request, 'quiz/quiz_result.html', {
        'quiz': quiz,
        'total': total_questions,
        'answered': answers.count(),
        'correct': correct,
        'percent': quiz.score,
    })



# 🔹 LOGIN sahifasi
def user_login(request):
    if request.user.is_authenticated:
        return redirect('quiztype_list')

    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            messages.success(request, f"Xush kelibsiz, {user.username}!")
            return redirect('quiztype_list')
        else:
            messages.error(request, "Login yoki parol noto‘g‘ri!")

    return render(request, 'quiz/login.html')


# 🔹 LOGOUT
def user_logout(request):
    logout(request)
    messages.info(request, "Tizimdan chiqdingiz.")
    return redirect('login')


# 🔹 SIGNUP (ro‘yxatdan o‘tish)
def user_signup(request):
    if request.user.is_authenticated:
        return redirect('quiztype_list')

    if request.method == 'POST':
        username = request.POST.get('username')
        first_name = request.POST.get('firstname')
        last_name = request.POST.get('lastname')
        email = request.POST.get('email')
        password1 = request.POST.get('password1')
        password2 = request.POST.get('password2')

        if password1 != password2:
            messages.error(request, "Parollar bir xil emas!")
            return redirect('signup')

        if User.objects.filter(username=username).exists():
            messages.error(request, "Bu foydalanuvchi nomi allaqachon mavjud.")
            return redirect('signup')

        user = User.objects.create_user(username=username, first_name=first_name, last_name=last_name, email=email, password=password1)
        login(request, user)
        messages.success(request, "Muvaffaqiyatli ro‘yxatdan o‘tdingiz!")
        return redirect('quiztype_list')

    return render(request, 'quiz/signup.html')

def result_users(request, test):
    user = request.user

    # Foydalanuvchi javoblarini savollar bilan birgalikda olish
    my_tests = (
        AnswerUsers.objects
        .filter(user=user, generate_quiz=test)
        .select_related('question', 'answer')
        .order_by('-id')
    )

    # Yuqorida select_related ishlatilgani uchun bu joyda alohida so‘rov kerak emas
    question_ids = my_tests.values_list('question_id', flat=True)
    answer = Answer.objects.filter(question_id__in=question_ids)

    context = {
        'user': user,
        'my_tests': my_tests,
        'answer': answer,
    }
    return render(request, 'quiz/result_users.html', context)